"""
Multimodal document processor: handles PDFs (text, tables, charts/images), 
DOCX, XLSX, CSV, and scanned images via OCR.
"""
import os
import io
import base64
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
import hashlib

from PIL import Image
import pytesseract
from pypdf import PdfReader
import pandas as pd

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {
    ".pdf", ".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".gif",
    ".docx", ".xlsx", ".csv", ".txt"
}


def get_file_hash(filepath: str) -> str:
    """Compute MD5 hash of file for dedup."""
    h = hashlib.md5()
    with open(filepath, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


def image_to_base64(image: Image.Image, max_size: Tuple[int, int] = (512, 512)) -> str:
    """Resize and encode a PIL image to base64."""
    image.thumbnail(max_size, Image.LANCZOS)
    buf = io.BytesIO()
    image.save(buf, format="PNG")
    return base64.b64encode(buf.getvalue()).decode("utf-8")


def ocr_image(image: Image.Image) -> str:
    """Run Tesseract OCR on a PIL image."""
    try:
        text = pytesseract.image_to_string(image, config="--oem 3 --psm 3")
        return text.strip()
    except Exception as e:
        logger.warning(f"OCR failed: {e}")
        return ""


def extract_pdf(filepath: str) -> List[Dict[str, Any]]:
    """
    Extract content from PDF:
    - Text pages → text chunks
    - Pages with embedded images → OCR + base64 stored in metadata
    - Tables detected via simple heuristic (pipe/tab-separated lines)
    Returns list of chunk dicts: {text, metadata}
    """
    chunks = []
    reader = PdfReader(filepath)
    filename = Path(filepath).name

    for page_num, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        
        # Detect table-like content
        lines = page_text.split("\n")
        table_lines = [l for l in lines if l.count("|") > 2 or l.count("\t") > 2]
        has_table = len(table_lines) > 3

        chunk_meta = {
            "source": filename,
            "page": page_num,
            "type": "table" if has_table else "text",
            "file_hash": get_file_hash(filepath),
        }

        if page_text.strip():
            chunks.append({
                "text": f"[Source: {filename}, Page {page_num}]\n{page_text.strip()}",
                "metadata": chunk_meta,
            })

    return chunks


def extract_image(filepath: str) -> List[Dict[str, Any]]:
    """OCR a standalone image file."""
    filename = Path(filepath).name
    pil_img = Image.open(filepath).convert("RGB")
    ocr_text = ocr_image(pil_img)
    # image_b64 intentionally omitted — not needed for vector retrieval
    return [{
        "text": f"[Source: {filename}]\n{ocr_text if ocr_text else '[Image with no detectable text]'}",
        "metadata": {
            "source": filename,
            "type": "image",
            "file_hash": get_file_hash(filepath),
        },
    }]


def extract_docx(filepath: str) -> List[Dict[str, Any]]:
    """Extract text and tables from DOCX."""
    from docx import Document
    filename = Path(filepath).name
    doc = Document(filepath)
    chunks = []
    file_hash = get_file_hash(filepath)

    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if full_text:
        chunks.append({
            "text": f"[Source: {filename}]\n{full_text}",
            "metadata": {"source": filename, "type": "text", "file_hash": file_hash},
        })

    for t_idx, table in enumerate(doc.tables):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        table_text = "\n".join(" | ".join(row) for row in rows)
        if table_text.strip():
            chunks.append({
                "text": f"[Source: {filename}, Table {t_idx+1}]\n{table_text}",
                "metadata": {"source": filename, "type": "table", "table_index": t_idx, "file_hash": file_hash},
            })
    return chunks


def extract_xlsx(filepath: str) -> List[Dict[str, Any]]:
    """Extract all sheets from XLSX as text."""
    filename = Path(filepath).name
    chunks = []
    file_hash = get_file_hash(filepath)
    xf = pd.ExcelFile(filepath)
    for sheet in xf.sheet_names:
        df = pd.read_excel(filepath, sheet_name=sheet)
        text = df.to_string(index=False)
        chunks.append({
            "text": f"[Source: {filename}, Sheet: {sheet}]\n{text}",
            "metadata": {"source": filename, "type": "table", "sheet": sheet, "file_hash": file_hash},
        })
    return chunks


def extract_csv(filepath: str) -> List[Dict[str, Any]]:
    filename = Path(filepath).name
    df = pd.read_csv(filepath)
    text = df.to_string(index=False)
    return [{
        "text": f"[Source: {filename}]\n{text}",
        "metadata": {"source": filename, "type": "table", "file_hash": get_file_hash(filepath)},
    }]


def extract_txt(filepath: str) -> List[Dict[str, Any]]:
    filename = Path(filepath).name
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    return [{
        "text": f"[Source: {filename}]\n{text}",
        "metadata": {"source": filename, "type": "text", "file_hash": get_file_hash(filepath)},
    }]


def process_document(filepath: str) -> List[Dict[str, Any]]:
    """Route file to the correct extractor."""
    ext = Path(filepath).suffix.lower()
    if ext == ".pdf":
        return extract_pdf(filepath)
    elif ext in {".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".gif"}:
        return extract_image(filepath)
    elif ext == ".docx":
        return extract_docx(filepath)
    elif ext == ".xlsx":
        return extract_xlsx(filepath)
    elif ext == ".csv":
        return extract_csv(filepath)
    elif ext == ".txt":
        return extract_txt(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def chunk_text(text: str, chunk_size: int = 800, overlap: int = 150) -> List[str]:
    """Split long text into overlapping chunks."""
    if len(text) <= chunk_size:
        return [text]
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunks.append(text[start:end])
        if end == len(text):
            break
        start += chunk_size - overlap
    return chunks


def ocr_text_to_dataframe(text: str):
    """Parse space-separated OCR table text into a DataFrame. Returns None if not table-like."""
    import re as _re
    from collections import Counter as _Counter

    _date_re = _re.compile(r'\d{4}-\d{2}-\d{2}|\d{1,2}/\d{1,2}/\d{4}')
    _num_re = _re.compile(r'^-?\d[\d,.]*$')
    _pipe_re = _re.compile(r'^\|+$')

    def _tokenize(line):
        return [t for t in line.split() if not _pipe_re.match(t)]

    def _is_header_candidate(tokens):
        if len(tokens) < 2:
            return False
        # Skip lines where every token is ≤2 chars — likely spreadsheet column letters
        if all(len(t.strip('._|')) <= 2 for t in tokens):
            return False
        if bool(_date_re.search(' '.join(tokens))):
            return False
        numeric = sum(1 for t in tokens if _num_re.match(t))
        alpha = sum(1 for t in tokens if _re.match(r'^[a-zA-Z#_]+$', t))
        is_leading_rownum = numeric == 1 and bool(_num_re.match(tokens[0].rstrip(',')))
        return alpha >= len(tokens) * 0.5 and (numeric == 0 or is_leading_rownum)

    def _merge_to_n(tokens, n_target):
        # Remove | artifacts from each token
        tokens = [t.replace('|', '') for t in tokens]
        tokens = [t for t in tokens if t]
        # Pass 1: tokens ending with '.' (e.g. "Rep.") merge into preceding token
        merged = []
        for t in tokens:
            if merged and t.endswith('.'):
                merged[-1] = merged[-1] + '_' + t.rstrip('.')
            else:
                merged.append(t)
        # Pass 2: tokens ending with '_' (OCR cell-border artifact) merge into preceding
        # e.g. "Unit" + "Price_" → "Unit_Price"
        merged2 = []
        for t in merged:
            if merged2 and t.endswith('_'):
                merged2[-1] = merged2[-1] + '_' + t.rstrip('_')
            else:
                merged2.append(t)
        # Pass 3: if still over target, merge the shortest adjacent pair
        while len(merged2) > n_target:
            best_i = min(range(len(merged2) - 1),
                         key=lambda i: len(merged2[i]) + len(merged2[i + 1]))
            merged2[best_i] = merged2[best_i] + '_' + merged2[best_i + 1]
            merged2.pop(best_i + 1)
        return merged2

    lines = [l.strip() for l in text.split('\n') if l.strip()]
    data_lines = [l for l in lines if not l.startswith('[Source:')]
    if len(data_lines) < 3:
        return None

    # Collect header candidates from first 15 lines
    candidates = []
    for i, line in enumerate(data_lines[:15]):
        tokens = _tokenize(line)
        if _is_header_candidate(tokens):
            candidates.append((i, tokens))

    if not candidates:
        return None

    # Score each candidate: determine expected column count from data-row token mode,
    # then count how many rows fall within ±2 of that count.
    best_idx = None
    best_score = -1
    best_skip_first = False
    best_raw_tokens = None
    best_n_data_cols = 0

    for cand_i, cand_tokens in candidates:
        raw = list(cand_tokens)
        skip_first = bool(_num_re.match(raw[0].rstrip(',')))
        if skip_first:
            raw = raw[1:]
        row_counts = []
        for line in data_lines[cand_i + 1:]:
            rtoks = _tokenize(line)
            if not rtoks or len(rtoks) < 2:
                continue
            if skip_first and rtoks[0][:1].isdigit():
                rtoks = rtoks[1:]
            row_counts.append(len(rtoks))
        if not row_counts:
            continue
        n_data_cols = _Counter(row_counts).most_common(1)[0][0]
        # Skip headers with fewer tokens than data columns — can't represent all columns
        if len(raw) < n_data_cols:
            continue
        score = sum(1 for c in row_counts if abs(c - n_data_cols) <= 2)
        if score > best_score:
            best_score = score
            best_idx = cand_i
            best_skip_first = skip_first
            best_raw_tokens = raw
            best_n_data_cols = n_data_cols

    if best_idx is None or best_score < 2:
        return None

    merged_headers = _merge_to_n(best_raw_tokens, best_n_data_cols)
    n_cols = len(merged_headers)

    # Dedupe column names
    seen: dict = {}
    final_headers = []
    for h in merged_headers:
        if h in seen:
            seen[h] += 1
            final_headers.append(f"{h}_{seen[h]}")
        else:
            seen[h] = 0
            final_headers.append(h)

    rows = []
    for line in data_lines[best_idx + 1:]:
        tokens = _tokenize(line)
        if not tokens or len(tokens) < 2:
            continue
        tokens = [t.rstrip(',') for t in tokens]
        if best_skip_first and tokens[0][:1].isdigit():
            tokens = tokens[1:]
        if len(tokens) > n_cols:
            row = tokens[:n_cols - 1] + [' '.join(tokens[n_cols - 1:])]
        else:
            row = tokens + [''] * (n_cols - len(tokens))
        # Skip mostly-empty rows (footer noise)
        if row.count('') >= max(1, n_cols // 2):
            continue
        rows.append(row)

    if len(rows) < 2:
        return None

    df = pd.DataFrame(rows, columns=final_headers)
    for col in df.columns:
        series = (df[col].str.replace(',', '', regex=False)
                         .str.replace('$', '', regex=False)
                         .str.replace('(', '-', regex=False)
                         .str.replace(')', '', regex=False))
        numeric = pd.to_numeric(series, errors='coerce')
        if numeric.notna().sum() > len(df) * 0.5:
            df[col] = numeric
            continue
        try:
            dates = pd.to_datetime(df[col], format='mixed', errors='coerce')
            if dates.notna().sum() > len(df) * 0.5:
                df[col] = dates
        except Exception:
            pass
    return df


def _img_table_is_useful(df) -> bool:
    """Return True if the DataFrame has usable table data.

    Accepts numeric tables AND text-only lookup tables (e.g. Sales Rep → Region).
    Rejects OCR false-positives: tables with <2 rows/cols, mostly-empty columns,
    or columns with only one distinct value.
    """
    if len(df) < 2 or len(df.columns) < 2:
        return False

    # Reject Excel column-header artifacts: first row is mostly None with single-char values (A, B, C...)
    first_row = df.iloc[0].astype(str).str.strip()
    none_mask = first_row.isin(["None", "nan", ""])
    if none_mask.mean() > 0.55:
        non_none = first_row[~none_mask]
        if not non_none.empty and (non_none.str.len() == 1).any():
            return False

    # Rule 2: any cell > 100 chars means prose/resume text or garbage — reject regardless of numeric columns
    all_vals = df.astype(str).values.flatten()
    if any(len(v.strip()) > 100 for v in all_vals if v.strip() not in ("nan", "None", "")):
        return False

    for col in df.columns:
        if pd.api.types.is_numeric_dtype(df[col]):
            return True
        parsed = pd.to_numeric(
            df[col].astype(str).str.replace(r"[$,\s%]", "", regex=True),
            errors="coerce",
        )
        if parsed.notna().mean() > 0.35:
            return True
    # Text-only: valid if every column has ≥50% non-empty values and ≥2 distinct values
    for col in df.columns:
        vals = df[col].astype(str).str.strip()
        non_empty = vals[(vals.str.len() > 0) & ~vals.isin(["nan", "None", ""])]
        if len(non_empty) / max(len(df), 1) < 0.5:
            return False
        if non_empty.nunique() < 2:
            return False
    return True


def _should_merge(bbox_a, bbox_b, y_gap_px: int = 50, x1_px: int = 20, x2_pct: float = 0.15) -> bool:
    """True when two bounding boxes represent vertically-adjacent segments of the same table."""
    y_gap = bbox_b.y1 - bbox_a.y2
    x_ref = max(bbox_a.x2, bbox_b.x2)
    return (
        y_gap < y_gap_px
        and abs(bbox_a.x1 - bbox_b.x1) <= x1_px
        and abs(bbox_a.x2 - bbox_b.x2) <= x_ref * x2_pct
    )


def _concat_dfs(base: "pd.DataFrame", extra: "pd.DataFrame") -> "pd.DataFrame":
    """Concatenate two DataFrames, aligning column count to base."""
    ncols = len(base.columns)
    n = len(extra.columns)
    extra = extra.copy()
    if n < ncols:
        for k in range(n, ncols):
            extra[f"_pad_{k}"] = ""
    elif n > ncols:
        extra = extra.iloc[:, :ncols]
    extra.columns = base.columns
    return pd.concat([base, extra], ignore_index=True)


def _merge_adjacent_img_tables(tables: list) -> list:
    """Group vertically-adjacent table segments with similar x-spans and merge each group.

    Works for any number of consecutive segments (not just pairs), so a table
    split into 3+ bounding-box regions is still returned as one DataFrame.

    Two segments are considered part of the same group when:
      - vertical gap < 50 px
      - left edges (x1) within 20 px
      - right edges (x2) within 15 % of the wider span
    """
    if len(tables) <= 1:
        return [t[0] for t in tables]

    tables = sorted(tables, key=lambda t: t[1].y1)

    # Build groups of consecutive adjacent segments
    groups: list[list] = [[tables[0]]]
    for df, bbox in tables[1:]:
        _, prev_bbox = groups[-1][-1]
        if _should_merge(prev_bbox, bbox):
            groups[-1].append((df, bbox))
        else:
            groups.append([(df, bbox)])

    result = []
    for group in groups:
        if len(group) == 1:
            result.append(group[0][0])
        else:
            merged = group[0][0]
            for df, _ in group[1:]:
                merged = _concat_dfs(merged, df)
            result.append(merged)
    return result


def _img2table_dfs(filepath: str) -> list:
    """Extract tables from a standalone image using img2table + Tesseract OCR.
    Returns a list of DataFrames (empty if no tables found or img2table unavailable).
    """
    try:
        from img2table.document import Image as _Img2Image
        from img2table.ocr import TesseractOCR as _TesseractOCR
        doc = _Img2Image(src=filepath)
        extracted = doc.extract_tables(
            ocr=_TesseractOCR(),
            implicit_rows=False,
            implicit_columns=False,
            borderless_tables=True,
        ) or []
        useful = [
            (t.df, t.bbox)
            for t in extracted
            if t.df is not None and not t.df.empty and _img_table_is_useful(t.df)
        ]
        return _merge_adjacent_img_tables(useful)
    except ImportError:
        logger.warning("img2table not installed; image table extraction skipped")
        return []
    except Exception as e:
        logger.warning(f"img2table extraction failed for '{filepath}': {e}")
        return []


def extract_dataframes(filepath: str) -> list:
    """Extract tables as DataFrames from a document. Returns empty list if none found."""
    ext = Path(filepath).suffix.lower()
    dfs = []
    try:
        if ext == '.csv':
            df = pd.read_csv(filepath)
            if not df.empty:
                dfs.append(df)
        elif ext == '.xlsx':
            xf = pd.ExcelFile(filepath)
            for sheet in xf.sheet_names:
                df = pd.read_excel(filepath, sheet_name=sheet)
                if not df.empty:
                    dfs.append(df)
        elif ext == '.docx':
            from docx import Document
            doc = Document(filepath)
            for table in doc.tables:
                rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                if len(rows) > 1:
                    df = pd.DataFrame(rows[1:], columns=rows[0])
                    if not df.empty:
                        dfs.append(df)
        elif ext == '.pdf':
            reader = PdfReader(filepath)
            for page in reader.pages:
                page_text = page.extract_text() or ''
                df = ocr_text_to_dataframe(page_text)
                if df is not None:
                    dfs.append(df)
        elif ext in {'.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif'}:
            dfs.extend(_img2table_dfs(filepath))
    except Exception as e:
        logger.warning(f"Table extraction failed for {filepath}: {e}")
    return dfs


def has_tables(filepath: str) -> bool:
    """Return True if the file type can contain extractable tables.

    CSV/XLSX are inherently tabular. DOCX is inspected for Word table objects.
    PDFs and images defer to extract_dataframes() as the authority — img2table
    runs on images so only files with real table structure produce DataFrames.
    """
    ext = Path(filepath).suffix.lower()
    if ext == '.txt':
        return False
    if ext in {'.csv', '.xlsx', '.xls'}:
        return True
    if ext == '.docx':
        try:
            from docx import Document as _DocxDoc
            return bool(_DocxDoc(filepath).tables)
        except Exception:
            return False
    if ext in {'.pdf', '.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif'}:
        return True  # extract_dataframes() is the authority
    return False


def extract_images(filepath: str) -> list:
    """
    Extract images from a document. Returns list of (page, img_idx, PIL.Image).
    - PDF: embedded images from every page (width/height >= 100px)
    - Standalone image files: the file itself as page=1, img_idx=0
    Other file types return an empty list.
    """
    ext = Path(filepath).suffix.lower()
    results = []
    if ext == ".pdf":
        reader = PdfReader(filepath)
        for page_num, page in enumerate(reader.pages, start=1):
            try:
                if not hasattr(page, "images") or not page.images:
                    continue
                for img_idx, img_obj in enumerate(page.images):
                    try:
                        pil_img = Image.open(io.BytesIO(img_obj.data)).convert("RGB")
                        if pil_img.width < 100 or pil_img.height < 100:
                            continue
                        results.append((page_num, img_idx, pil_img))
                    except Exception as e:
                        logger.debug(f"Skipping image p{page_num}[{img_idx}]: {e}")
            except Exception as e:
                logger.debug(f"Image extraction error on page {page_num}: {e}")
    elif ext == ".docx":
        try:
            from docx import Document as _Document
            doc = _Document(filepath)
            img_idx = 0
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    try:
                        img_data = rel.target_part.blob
                        pil_img = Image.open(io.BytesIO(img_data)).convert("RGB")
                        if pil_img.width < 100 or pil_img.height < 100:
                            continue
                        results.append((1, img_idx, pil_img))
                        img_idx += 1
                    except Exception as e:
                        logger.debug(f"Skipping DOCX image: {e}")
        except Exception as e:
            logger.warning(f"DOCX image extraction failed for {filepath}: {e}")
    elif ext in {".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".gif"}:
        try:
            pil_img = Image.open(filepath).convert("RGB")
            results.append((1, 0, pil_img))
        except Exception as e:
            logger.warning(f"Failed to open image file {filepath}: {e}")
    return results


def images_to_chunks(filepath: str, images: list) -> List[Dict[str, Any]]:
    """OCR extracted images and return chunked text records for the vector store.

    images: list of (page, img_idx, PIL.Image) as returned by extract_images().
    Only images that produce non-empty OCR text are included.
    """
    filename = Path(filepath).name
    chunks = []
    for page, img_idx, pil_img in images:
        ocr_text = ocr_image(pil_img)
        if not ocr_text:
            continue
        header = f"[Source: {filename}, Page {page}, Image {img_idx}]"
        for i, sub in enumerate(chunk_text(f"{header}\n{ocr_text}")):
            chunks.append({
                "text": sub,
                "metadata": {
                    "source": filename,
                    "page": page,
                    "image_index": img_idx,
                    "type": "image",
                    "chunk_index": i,
                },
            })
    return chunks


def process_document_chunked(filepath: str) -> List[Dict[str, Any]]:
    """Process a document and chunk large text blocks."""
    raw_chunks = process_document(filepath)
    final_chunks = []
    for chunk in raw_chunks:
        text = chunk["text"]
        meta = chunk["metadata"]
        sub_texts = chunk_text(text)
        for i, sub in enumerate(sub_texts):
            final_chunks.append({
                "text": sub,
                "metadata": {**meta, "chunk_index": i},
            })
    return final_chunks
